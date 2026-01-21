"""
Resume Parser Service
Extracts text from PDF and DOC/DOCX resume files
"""

import os
from typing import Optional
from pathlib import Path

class ResumeParser:
    """Parses resume files (PDF, DOC, DOCX) and extracts text"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.doc', '.docx']
    
    def parse(self, file_path: str) -> str:
        """
        Parse resume file and extract text content
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Extracted text content
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported: {self.supported_formats}")
        
        if file_ext == '.pdf':
            return self._parse_pdf(file_path)
        elif file_ext in ['.doc', '.docx']:
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Parser not implemented for: {file_ext}")
    
    def _parse_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install PyPDF2")
        
        text_content = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
            
            return '\n\n'.join(text_content)
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")
    
    def _parse_docx(self, file_path: str) -> str:
        """Extract text from DOC/DOCX file"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOC/DOCX parsing. Install with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            return '\n\n'.join(text_content)
        except Exception as e:
            raise ValueError(f"Error parsing DOC/DOCX: {str(e)}")
    
    def is_supported(self, filename: str) -> bool:
        """Check if file format is supported"""
        file_ext = Path(filename).suffix.lower()
        return file_ext in self.supported_formats

resume_parser = ResumeParser()
